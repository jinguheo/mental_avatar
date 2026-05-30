"""Word(.docx) 파서 — python-docx 사용"""
import os
from docx import Document

CHUNK_SIZE = 1500


def parse_docx(file_path: str) -> list[dict]:
    doc = Document(file_path)
    base_title = os.path.splitext(os.path.basename(file_path))[0]

    sections = []
    current_heading = base_title
    current_text = ""

    for para in doc.paragraphs:
        style = para.style.name
        text = para.text.strip()
        if not text:
            continue

        if style.startswith("Heading"):
            # 이전 섹션 저장
            if current_text.strip():
                sections.append({"heading": current_heading, "content": current_text.strip()})
            current_heading = text
            current_text = ""
        else:
            current_text += text + "\n"

    if current_text.strip():
        sections.append({"heading": current_heading, "content": current_text.strip()})

    # 섹션을 청크로 변환
    chunks = []
    for idx, sec in enumerate(sections):
        chunks.append({
            "title": sec["heading"],
            "content": sec["content"][:CHUNK_SIZE],
            "meta": {"chunk_index": idx, "source_type": "docx", "file_path": file_path}
        })

    # 표 내용도 추출
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            chunks.append({
                "title": f"{base_title} [표]",
                "content": "\n".join(rows),
                "meta": {"source_type": "docx_table", "file_path": file_path}
            })

    return chunks if chunks else [{"title": base_title, "content": "", "meta": {"source_type": "docx", "file_path": file_path}}]
